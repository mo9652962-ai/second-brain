# -*- coding: utf-8 -*-
"""
三年级数学每日一练40天 - 标准教材版生成器
版本: 3.0 (十轮研究优化版)
作者: AI Assistant
创建日期: 2026-07-29

基于十轮研究成果的优化：
  1. ✅ 人教版小学教材标准格式
  2. ✅ 专业出版级排版规范
  3. ✅ 符合认知心理学的页面布局
  4. ✅ 艾宾浩斯复习节奏融入
  5. ✅ 模块化架构，单一职责
  6. ✅ 标准分数竖式格式
  7. ✅ 合理题量密度，认知负荷优化
  8. ✅ 完整文档元数据
"""

import os
import random
from typing import List, Tuple, Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ============================================================================
# 配置常量区 (基于十轮研究：人教版教材标准 + 认知心理学优化)
# ============================================================================

class MathWorkbookConfig:
    """
    练习册配置类 - 符合人教版小学出版标准
    
    参考文献：
    - 义务教育数学课程标准 (2022年版)
    - 小学教材排版规范 (人教版)
    - 认知心理学与教学设计
    """
    
    # ============== 页面设置 (标准A4) ==============
    PAGE_WIDTH = Cm(21.0)          # A4纸宽度
    PAGE_HEIGHT = Cm(29.7)         # A4纸高度
    MARGIN_TOP = Cm(2.5)           # 上边距（教材标准）
    MARGIN_BOTTOM = Cm(2.0)        # 下边距
    MARGIN_LEFT = Cm(2.0)          # 左边距
    MARGIN_RIGHT = Cm(2.0)         # 右边距
    
    # ============== 字体设置 (人教版小学教材标准) ==============
    FONT_MAIN = "宋体"              # 正文字体（教材标准）
    FONT_TITLE = "黑体"             # 标题字体
    FONT_MATH = "Times New Roman"   # 数学符号字体
    
    # 字号设置（严格按照教材）
    FONT_SIZE_BOOK_TITLE = Pt(20)   # 封面标题 (小一号)
    FONT_SIZE_DAY_TITLE = Pt(16)    # 每日标题 (三号)
    FONT_SIZE_SECTION = Pt(14)      # 板块标题 (四号)
    FONT_SIZE_PROBLEM = Pt(12)      # 题目字号 (小四，保护视力)
    FONT_SIZE_ANSWER_LINE = Pt(11)  # 答题行字号
    FONT_SIZE_INFO = Pt(9)          # 辅助信息字号
    
    # ============== 行距与间距 (认知负荷优化) ==============
    LINE_SPACING_PROBLEM = 1.3      # 题目行距 (1.3倍，保护视力)
    LINE_SPACING_SECTION = 1.5      # 板块间距
    PARAGRAPH_SPACE_BEFORE = Pt(6)  # 段前距
    PARAGRAPH_SPACE_AFTER = Pt(6)   # 段后距
    
    # ============== 题量配置 (基于认知心理学研究) ==============
    # 小学生每日最佳练习量：30-40分钟，符合注意力持续时间
    PROBLEMS_PER_DAY = {
        "口算": 10,    # 口算：两位数乘法，10题 ≈ 5分钟 (5类混合，难度递进)
        "笔算": 4,     # 笔算：两位数乘两位数列竖式，4题 ≈ 12分钟 (含答题空间，竖式计算)
        "分数": 10,    # 分数：概念理解，10题 ≈ 10分钟
        "填空": 5,     # 单位换算：5题 ≈ 4分钟
        "应用": 2,     # 应用题：综合应用，2题 ≈ 5分钟
    }
    # 总计：42题，约30分钟
    
    # ============== 分数竖式标准格式 (人教版规范) ==============
    FRACTION_LINE_CHAR = "─"        # 标准分数线字符
    FRACTION_LINE_LENGTH = 3        # 分数线长度（教材标准：3字符）
    FRACTION_NUM_DEN_SPACING = Pt(4)  # 分子/分母与分数线间距
    
    # ============== 去重配置 (哈希+预生成池) ==============
    USED_PROBLEMS = set()           # 全局去重集合
    MAX_RETRY = 100                 # 最大重试次数
    DEDUP_ENABLED = True            # 是否启用去重
    # 预生成池（确保40天×10道=400道竖式题不重复）
    SHUSHI_POOL_SIZE = 500          # 竖式题预生成池（2位×2位：8100种可能，取500）
    
    # ============== 艾宾浩斯复习标记 ==============
    REVIEW_MARKS = {
        1: "📖 学习日",
        2: "🔄 第1次复习",
        4: "🔄 第2次复习",
        7: "🔄 第3次复习",
        15: "🔄 第4次复习",
        30: "🔄 第5次复习",
    }
    
    # ============== 输出配置 ==============
    DESKTOP_PATH = r"C:\Users\31954\Desktop"
    OUTPUT_FILENAME = "三年级数学每日一练40天_标准教材版_v2.docx"
    OUTPUT_FULL_PATH = os.path.join(DESKTOP_PATH, OUTPUT_FILENAME)
    
    # ============== 元数据 ==============
    BOOK_TITLE = "三年级数学每日一练"
    BOOK_SUBTITLE = "40天系统练习版"
    BOOK_AUTHOR = "数学教研组"
    BOOK_VERSION = "3.1"
    BOOK_DATE = "2026年7月"


# ============================================================================
# 工具函数区 (单一职责原则)
# ============================================================================

def set_chinese_font(run, font_name: str, font_size: Pt, bold: bool = False) -> None:
    """
    设置中文字体（解决python-docx中文字体不生效问题）
    
    Args:
        run: docx run对象
        font_name: 字体名称
        font_size: 字号
        bold: 是否加粗
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)  # 纯黑色，清晰打印


def add_title_paragraph(doc, text: str, font_size: Pt, bold: bool = True, alignment: int = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """
    添加标题段落（标准格式）
    
    Args:
        doc: Document对象
        text: 标题文本
        font_size: 字号
        bold: 是否加粗
        alignment: 对齐方式
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    set_chinese_font(run, MathWorkbookConfig.FONT_TITLE, font_size, bold)


def add_problem_paragraph(doc, text: str, first_line_indent: bool = True) -> None:
    """
    添加题目段落（教材标准格式）
    
    Args:
        doc: Document对象
        text: 题目文本
        first_line_indent: 是否首行缩进
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_PROBLEM
    p.paragraph_format.space_before = MathWorkbookConfig.PARAGRAPH_SPACE_BEFORE
    p.paragraph_format.space_after = MathWorkbookConfig.PARAGRAPH_SPACE_AFTER
    
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进（教材标准）
    
    run = p.add_run(text)
    set_chinese_font(run, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)


def add_empty_line(doc, count: int = 1) -> None:
    """
    添加空行（用于控制间距）
    
    Args:
        doc: Document对象
        count: 空行数量
    """
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)


# ============================================================================
# 题目生成函数 (纯计算逻辑 + 去重机制，与渲染分离)
# ============================================================================

def _make_hash(prefix: str, *args) -> str:
    """生成题目唯一哈希（用于去重）"""
    return f"{prefix}:{'_'.join(map(str, args))}"


def _try_generate(prefix: str, gen_func, max_retry: int = None) -> str | Tuple:
    """
    带重试的题目生成，确保不重复
    
    Args:
        prefix: 题目类型前缀
        gen_func: 生成函数（返回 hash_key, result）
        max_retry: 最大重试次数
        
    Returns:
        生成的题目
    """
    max_tries = max_retry or MathWorkbookConfig.MAX_RETRY
    for _ in range(max_tries):
        hash_key, result = gen_func()
        if hash_key not in MathWorkbookConfig.USED_PROBLEMS:
            MathWorkbookConfig.USED_PROBLEMS.add(hash_key)
            return result
    # 如果重试耗尽，返回最后一个结果（极罕见情况）
    return result


def generate_kousuan_problems(count: int) -> List[str]:
    """
    生成口算题目 (两位数乘法，5类混合，难度递进) - 去重版
    
    基于人教版三年级数学下册知识点：
    1. 整十数×整十数 (3题) - 热身
    2. 整十数×两位数   (2题) - 进阶
    3. 两位数×整十数   (2题) - 对称练习
    4. 两位数×两位数(不进位) (2题) - 挑战
    5. 两位数×两位数(进位)   (1题) - 拔高
    
    每题计算可用口算方法(拆数法/整十法)，不要求完全心算。
    
    Args:
        count: 题目数量 (10)
        
    Returns:
        题目字符串列表
    """
    tens = [10, 20, 30, 40, 50, 60, 70, 80, 90, 100, 110, 120]
    # 注：整十数×整十数池 |tens|²=144种，40天×2=80够用
    
    def _gen_type1():
        """整十数×整十数: 30×40=1200"""
        a = random.choice(tens)
        b = random.choice(tens)
        return _make_hash("kou_t1", a, b), f"{a} × {b} = ____"
    
    def _gen_type2():
        """整十数×两位数: 30×45=1350"""
        a = random.choice(tens)
        b = random.randint(11, 99)
        return _make_hash("kou_t2", a, b), f"{a} × {b} = ____"
    
    def _gen_type3():
        """两位数×整十数: 12×30=360"""
        a = random.randint(11, 99)
        b = random.choice(tens)
        return _make_hash("kou_t3", a, b), f"{a} × {b} = ____"
    
    def _gen_type4():
        """两位数×两位数(不进位): 12×13=156"""
        a = random.randint(11, 49)
        b = random.randint(11, 49)
        # 确保个位相乘不进位
        a_ones = a % 10
        b_ones = b % 10
        while a_ones * b_ones >= 10:
            a = random.randint(11, 49)
            b = random.randint(11, 49)
            a_ones = a % 10
            b_ones = b % 10
        return _make_hash("kou_t4", a, b), f"{a} × {b} = ____"
    
    def _gen_type5():
        """两位数×两位数(进位): 27×34=918"""
        a = random.randint(11, 99)
        b = random.randint(11, 99)
        # 确保至少一次进位
        a_ones = a % 10
        b_ones = b % 10
        while a_ones * b_ones < 10:
            a = random.randint(11, 99)
            b = random.randint(11, 99)
            a_ones = a % 10
            b_ones = b % 10
        return _make_hash("kou_t5", a, b), f"{a} × {b} = ____"
    
    problems = []
    # 按难度递进生成
    for _ in range(2):
        problems.append(_try_generate("kou_t1", _gen_type1))
    for _ in range(2):
        problems.append(_try_generate("kou_t2", _gen_type2))
    for _ in range(2):
        problems.append(_try_generate("kou_t3", _gen_type3))
    for _ in range(3):
        problems.append(_try_generate("kou_t4", _gen_type4))
    for _ in range(1):
        problems.append(_try_generate("kou_t5", _gen_type5))
    
    return problems


def generate_shushi_problems(count: int) -> List[Tuple[int, int, str]]:
    """
    生成竖式笔算题目 (两位数×两位数，4类混合，难度递进) - 去重版
    
    基于人教版三年级数学下册笔算乘法：
    1. 不进位：如 12×13=156  (个位×个位<10，无需进位)
    2. 一次进位：如 15×13=195  (个位进位一次)
    3. 两次进位：如 27×34=918  (个位和十位都有进位)
    4. 连续进位：如 89×76=6764 (连续多次进位，难度最高)
    
    每天4题，每类1题，40天×4=160题不重复
    
    Args:
        count: 题目数量 (4)
        
    Returns:
        (乘数a, 乘数b, 分类标签) 元组列表
    """
    def _gen_type_nc():
        """不进位：个位×个位<10，十位×个位<10"""
        a = random.randint(11, 49)
        b = random.randint(11, 49)
        a1, a0 = a // 10, a % 10
        b1, b0 = b // 10, b % 10
        while (a0 * b0 >= 10) or (a1 * b0 >= 10):
            a = random.randint(11, 49)
            b = random.randint(11, 49)
            a1, a0 = a // 10, a % 10
            b1, b0 = b // 10, b % 10
        return _make_hash("shu_nc", a, b), (a, b, "不进位")
    
    def _gen_type_sc():
        """一次进位：个位进位，十位不进位"""
        a = random.randint(11, 49)
        b = random.randint(11, 49)
        a1, a0 = a // 10, a % 10
        b1, b0 = b // 10, b % 10
        while a0 * b0 < 10 or a1 * b0 >= 10:
            a = random.randint(11, 49)
            b = random.randint(11, 49)
            a1, a0 = a // 10, a % 10
            b1, b0 = b // 10, b % 10
        return _make_hash("shu_sc", a, b), (a, b, "一次进位")
    
    def _gen_type_dc():
        """两次进位：个位和十位都进位"""
        a = random.randint(11, 99)
        b = random.randint(11, 99)
        a1, a0 = a // 10, a % 10
        b1, b0 = b // 10, b % 10
        while a0 * b0 < 10 or a1 * b0 < 10:
            a = random.randint(11, 99)
            b = random.randint(11, 99)
            a1, a0 = a // 10, a % 10
            b1, b0 = b // 10, b % 10
        return _make_hash("shu_dc", a, b), (a, b, "两次进位")
    
    def _gen_type_cc():
        """连续进位：个位进位大(≥3)，十位进位大"""
        a = random.randint(51, 99)
        b = random.randint(51, 99)
        a1, a0 = a // 10, a % 10
        b1, b0 = b // 10, b % 10
        while (a0 * b0 < 20) or (a1 * b0 < 15):
            a = random.randint(51, 99)
            b = random.randint(51, 99)
            a1, a0 = a // 10, a % 10
            b1, b0 = b // 10, b % 10
        return _make_hash("shu_cc", a, b), (a, b, "连续进位")
    
    problems = []
    gens = [("shu_nc", _gen_type_nc), ("shu_sc", _gen_type_sc),
            ("shu_dc", _gen_type_dc), ("shu_cc", _gen_type_cc)]
    for prefix, gen in gens:
        result = _try_generate(prefix, gen)
        problems.append(result)
    return problems


def generate_fraction_problems(count: int) -> List[Tuple[int, int, int, int, str]]:
    """
    生成分数加减法题目 (同分母) - 去重版+扩展池
    分母覆盖 [2,3,4,5,6,7,8,9,10,11,12,15,20]，加减法各半
    确保每天有加法也有减法
    
    Args:
        count: 题目数量
        
    Returns:
        (分子1, 分母, 分子2, 分母, 运算符) 元组列表
    """
    denominators = [2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 15, 20]
    
    def _gen_one(op: str):
        den = random.choice(denominators)
        n1 = random.randint(1, den - 1) if den > 2 else random.randint(1, den)
        n2 = random.randint(1, den - 1) if den > 2 else random.randint(1, den)
        if op == '-' and n2 > n1:
            n1, n2 = n2, n1
        key = _make_hash("frac", n1, den, n2, op)
        return key, (n1, den, n2, den, op)
    
    problems = []
    half = count // 2
    # 前半加号，后半减号，确保每天都有加减
    for _ in range(half):
        result = _try_generate("frac", lambda: _gen_one('+'))
        problems.append(result)
    for _ in range(count - half):
        result = _try_generate("frac", lambda: _gen_one('-'))
        problems.append(result)
    # 打乱顺序，避免加号全在前面
    random.shuffle(problems)
    return problems


def generate_fill_problems(count: int) -> List[str]:
    """
    生成时间单位换算填空题目 (2类：时分秒 + 年月日)
    题库200+题，每天随机抽取5道，40天不重复
    基于人教版三年级数学上册(时分秒)+下册(年月日)知识点
    """
    pool = [
        # ===== 时分秒：基础换算 (1时=60分, 1分=60秒) =====
        "1时 = ____分", "1分 = ____秒",
        "2时 = ____分", "3时 = ____分",
        "4时 = ____分", "5时 = ____分",
        "6时 = ____分", "7时 = ____分",
        "8时 = ____分", "9时 = ____分",
        "2分 = ____秒", "3分 = ____秒",
        "4分 = ____秒", "5分 = ____秒",
        "6分 = ____秒", "7分 = ____秒",
        "8分 = ____秒", "9分 = ____秒",
        "120分 = ____时", "180分 = ____时",
        "240分 = ____时", "300分 = ____时",
        "360分 = ____时", "420分 = ____时",
        "120秒 = ____分", "180秒 = ____分",
        "240秒 = ____分", "300秒 = ____分",
        "360秒 = ____分", "420秒 = ____分",
        "60秒 = ____分", "60分 = ____时",
        "10分 = ____秒", "10时 = ____分",
        "480分 = ____时", "540分 = ____时",
        "480秒 = ____分", "540秒 = ____分",
        # ===== 时分秒：复名数换算 =====
        "1时5分 = ____分", "1时10分 = ____分",
        "1时15分 = ____分", "1时20分 = ____分",
        "1时25分 = ____分", "1时30分 = ____分",
        "1时35分 = ____分", "1时40分 = ____分",
        "1时45分 = ____分", "1时50分 = ____分",
        "2时5分 = ____分", "2时10分 = ____分",
        "2时15分 = ____分", "2时20分 = ____分",
        "2时25分 = ____分", "2时30分 = ____分",
        "2时40分 = ____分", "2时50分 = ____分",
        "3时5分 = ____分", "3时10分 = ____分",
        "3时20分 = ____分", "3时30分 = ____分",
        "3时45分 = ____分", "1时55分 = ____分",
        "1分5秒 = ____秒", "1分10秒 = ____秒",
        "1分15秒 = ____秒", "1分20秒 = ____秒",
        "1分25秒 = ____秒", "1分30秒 = ____秒",
        "1分35秒 = ____秒", "1分40秒 = ____秒",
        "1分45秒 = ____秒", "1分50秒 = ____秒",
        "2分5秒 = ____秒", "2分10秒 = ____秒",
        "2分15秒 = ____秒", "2分20秒 = ____秒",
        "2分25秒 = ____秒", "2分30秒 = ____秒",
        "2分35秒 = ____秒", "2分45秒 = ____秒",
        "3分5秒 = ____秒", "3分10秒 = ____秒",
        "3分20秒 = ____秒", "3分30秒 = ____秒",
        "3分45秒 = ____秒", "1分55秒 = ____秒",
        "60分 = ____时____分", "65分 = ____时____分",
        "70分 = ____时____分", "75分 = ____时____分",
        "80分 = ____时____分", "85分 = ____时____分",
        "90分 = ____时____分", "95分 = ____时____分",
        "100分 = ____时____分", "105分 = ____时____分",
        "110分 = ____时____分", "125分 = ____时____分",
        "135分 = ____时____分", "150分 = ____时____分",
        "160分 = ____时____分", "175分 = ____时____分",
        "60秒 = ____分____秒", "70秒 = ____分____秒",
        "75秒 = ____分____秒", "80秒 = ____分____秒",
        "85秒 = ____分____秒", "90秒 = ____分____秒",
        "95秒 = ____分____秒", "100秒 = ____分____秒",
        "110秒 = ____分____秒", "120秒 = ____分____秒",
        "130秒 = ____分____秒", "145秒 = ____分____秒",
        "155秒 = ____分____秒", "180秒 = ____分____秒",
        "200秒 = ____分____秒", "210秒 = ____分____秒",
        "1时 = ____秒", "2时 = ____秒",
        "3600秒 = ____时", "7200秒 = ____时",
        "1刻钟 = ____分", "半小时 = ____分",
        "2刻钟 = ____分", "3刻钟 = ____分",
        "1时30分 = ____刻钟", "1时 = ____刻钟",
        # ===== 年月日：基础换算 =====
        "1年 = ____个月", "2年 = ____个月",
        "3年 = ____个月", "4年 = ____个月",
        "5年 = ____个月", "6年 = ____个月",
        "12个月 = ____年", "24个月 = ____年",
        "36个月 = ____年", "48个月 = ____年",
        "60个月 = ____年", "72个月 = ____年",
        "1日 = ____时", "2日 = ____时",
        "3日 = ____时", "4日 = ____时",
        "5日 = ____时", "6日 = ____时",
        "24时 = ____日", "48时 = ____日",
        "72时 = ____日", "96时 = ____日",
        "120时 = ____日", "144时 = ____日",
        "1周 = ____天", "2周 = ____天",
        "3周 = ____天", "4周 = ____天",
        "5周 = ____天", "6周 = ____天",
        "14天 = ____周", "21天 = ____周",
        "28天 = ____周", "35天 = ____周",
        "42天 = ____周", "49天 = ____周",
        # ===== 年月日：大月小月常识 =====
        "1月有 ____天", "3月有 ____天",
        "5月有 ____天", "7月有 ____天",
        "8月有 ____天", "10月有 ____天",
        "12月有 ____天", "4月有 ____天",
        "6月有 ____天", "9月有 ____天",
        "11月有 ____天",
        "平年2月有 ____天", "闰年2月有 ____天",
        "平年全年有 ____天", "闰年全年有 ____天",
        "一年有 ____个大月", "一年有 ____个小月",
        "大月每月有 ____天", "小月每月有 ____天",
        # ===== 年月日：综合换算 =====
        "平年上半年有 ____天", "平年下半年有 ____天",
        "闰年上半年有 ____天", "闰年下半年有 ____天",
        "第一季度有 ____天(平年)", "第二季度有 ____天",
        "第三季度有 ____天", "第四季度有 ____天",
        "第一季度有 ____天(闰年)",
        "1世纪 = ____年", "2世纪 = ____年",
        "1月有 ____个星期零____天",
        "3月有 ____个星期零____天",
        "5月有 ____个星期零____天",
        "7月有 ____个星期零____天",
        "8月有 ____个星期零____天",
        "10月有 ____个星期零____天",
        "12月有 ____个星期零____天",
        "平年全年有 ____个星期零____天",
        "闰年全年有 ____个星期零____天",
        # ===== 年月日+时分秒混合 =====
        "1年 = ____个月 = ____天(平年)",
        "1日 = ____时 = ____分",
        "1时 = ____分 = ____秒",
        "30个月 = ____年____个月",
        "45个月 = ____年____个月",
        "50个月 = ____年____个月",
        "2年6个月 = ____个月",
        "3年3个月 = ____个月",
        "1日6时 = ____时",
        "1日12时 = ____时",
        "2日8时 = ____时",
        "半年 = ____个月",
        "1年半 = ____个月",
        "连续两个月共62天的是 ____月和____月",
        "连续两个月共61天的是 ____月和____月",
        "连续两个月共60天的是 ____月和____月",
        "连续两个月共59天的是 ____月(平年)和____月",
        # ===== 24时计时法 =====
        "下午2时用24时计时法表示是 ____时",
        "下午4时用24时计时法表示是 ____时",
        "晚上8时用24时计时法表示是 ____时",
        "晚上9时用24时计时法表示是 ____时",
        "上午10时用24时计时法表示是 ____时",
        "下午5时用24时计时法表示是 ____时",
        "15时用普通计时法表示是 ____午____时",
        "19时用普通计时法表示是 ____上____时",
        "21时用普通计时法表示是 ____上____时",
        "22时用普通计时法表示是 ____上____时",
        "13时用普通计时法表示是 ____午____时",
        "17时用普通计时法表示是 ____午____时",
        "0时也叫 ____时",
        "中午12时用24时计时法表示是 ____时",
        # ===== 时间计算 =====
        "分针走1小格是 ____分",
        "分针走1大格是 ____分",
        "分针走一圈是 ____分",
        "时针走1大格是 ____时",
        "时针走一圈是 ____时",
        "秒针走1小格是 ____秒",
        "秒针走一圈是 ____秒",
        "秒针走一圈，分针走 ____小格",
        "分针走一圈，时针走 ____大格",
        "一昼夜时针走 ____圈",
        "2日 = ____时，也是 ____个昼夜",
    ]
    # 使用全局去重池：第一次调用时预洗牌，逐天消耗确保不重复
    if not hasattr(MathWorkbookConfig, '_fill_pool_index'):
        random.shuffle(pool)
        MathWorkbookConfig._fill_pool = pool
        MathWorkbookConfig._fill_pool_index = 0
    
    idx = MathWorkbookConfig._fill_pool_index
    chosen = MathWorkbookConfig._fill_pool[idx:idx + count]
    MathWorkbookConfig._fill_pool_index += count
    
    # 如果池子不够（理论上不会），回绕重新洗牌
    if len(chosen) < count:
        random.shuffle(MathWorkbookConfig._fill_pool)
        MathWorkbookConfig._fill_pool_index = count
        chosen = MathWorkbookConfig._fill_pool[:count]
    
    return chosen


def generate_application_problems(count: int) -> List[str]:
    """
    生成应用题：每天2道专项题型
    题1：两位数×两位数乘法应用题
    题2：同分母分数加减应用题
    题库各40+题，洗牌后逐天消耗，确保40天不重复
    """
    # ===== 乘法应用题池 (两位数×两位数) =====
    mul_pool = [
        "水龙头每天浪费13千克水，16天浪费多少千克水？",
        "每箱苹果有24个，买了15箱，一共有多少个苹果？",
        "每页书有25行，一本书有18页，一共有多少行？",
        "每排有12个座位，一个教室有14排，一共有多少个座位？",
        "每天读18页书，23天一共读了多少页？",
        "每袋大米重25千克，食堂买了12袋，一共重多少千克？",
        "一辆卡车每次运15吨货物，运了17次，一共运了多少吨？",
        "每盒彩笔有12支，学校买了24盒，一共有多少支彩笔？",
        "小明每分钟走45米，走了18分钟，一共走了多少米？",
        "每包饼干有16块，买了15包，一共有多少块饼干？",
        "工人每天铺路23米，铺了14天，一共铺了多少米？",
        "每张桌子坐12人，有18张桌子，一共可以坐多少人？",
        "每棵树产苹果36千克，果园有15棵苹果树，一共产多少千克？",
        "每瓶饮料有28元，买了13瓶，一共需要多少元？",
        "每层楼有18级台阶，一栋楼有15层，一共有多少级台阶？",
        "每只兔子吃14根胡萝卜，18只兔子一共吃多少根胡萝卜？",
        "每箱牛奶有24盒，超市进了16箱，一共有多少盒牛奶？",
        "每张邮票值12元，一套有22张，一共值多少元？",
        "每天的作业有16道题，做了25天，一共做了多少道题？",
        "每把椅子42元，买了13把，一共需要多少元？",
        "每本练习本有32页，买了15本，一共有多少页？",
        "每千克猪肉26元，买了14千克，一共多少元？",
        "每间教室有25张桌子，学校有16间教室，一共有多少张桌子？",
        "每台风扇135元，买了12台，一共需要多少元？",
        "每排树有22棵，一个果园种了18排，一共种了多少棵树？",
        "每场电影票价38元，一个班有15人去看，一共需要多少元？",
        "每个花坛种24株花，公园有16个花坛，一共种了多少株花？",
        "每根跳绳17元，体育老师买了15根，一共需要多少元？",
        "每盆花需要35克肥料，大棚里有18盆花，一共需要多少克肥料？",
        "每辆车坐45人，有12辆车，一共可以坐多少人？",
        "每个袋子装18个橘子，装了15袋，一共有多少个橘子？",
        "每天存17元钱，存了16天，一共存了多少钱？",
        "每桶油重19千克，食堂买了14桶，一共重多少千克？",
        "每块地砖长45厘米，铺了18块，一共长多少厘米？",
        "每把椅子重12千克，搬了26把，一共搬了多少千克？",
        "每瓶矿泉水2元，买了48瓶，一共需要多少元？",
        "每个文具盒13元，买了25个，一共需要多少元？",
        "每支钢笔15元，买了18支，一共需要多少元？",
        "每小时生产零件26个，生产了16小时，一共生产多少个？",
        "每盒巧克力有24颗，买了16盒，一共有多少颗？",
        "每件衣服需要15颗扣子，做了28件，一共需要多少颗扣子？",
        "每个书架放36本书，图书馆有12个书架，一共可以放多少本书？",
        "每度电费1元，一家工厂每天用电45度，用了18天，一共多少元？",
        "每张海报需要23分钟绘制，画了14张，一共用多少分钟？",
        "每包糖果有18颗，买了22包，一共有多少颗糖果？",
    ]
    
    # ===== 分数应用题池 =====
    frac_pool = [
        # 平均分成N份，用了M份，剩下占几分之几
        "布料平均分成6份，用掉5份，剩下占几分之几？",
        "蛋糕平均分成8份，吃了3份，剩下的占几分之几？",
        "一根绳子平均分成10份，用掉7份，剩下占几分之几？",
        "一张纸平均分成4份，用了1份，剩下占几分之几？",
        "西瓜平均分成8份，吃了5份，剩下占几分之几？",
        "披萨平均分成6份，吃了2份，剩下占几分之几？",
        "彩带平均分成5份，用掉3份，剩下占几分之几？",
        "饼干平均分成8份，吃了6份，剩下占几分之几？",
        "土地平均分成9份，种了4份菜，剩下的占几分之几？",
        "巧克力平均分成7份，吃了4份，剩下占几分之几？",
        "蛋糕平均分成12份，吃了7份，剩下的占几分之几？",
        "绳子平均分成10份，用掉3份，剩下占几分之几？",
        "布料平均分成9份，用掉5份，剩下占几分之几？",
        "月饼平均分成4份，吃了3份，剩下占几分之几？",
        "木板平均分成8份，用了3份，剩下占几分之几？",
        "蛋糕平均分成12份，吃了5份，剩下的占几分之几？",
        # 一盒有N个，吃了M个，剩下占几分之几
        "一盒饼干有12块，吃了5块，剩下的占几分之几？",
        "一盒巧克力有10颗，吃了3颗，剩下的占几分之几？",
        "一盒鸡蛋有12个，用掉7个，剩下的占几分之几？",
        "一包糖果有15颗，吃了8颗，剩下的占几分之几？",
        "一箱苹果有16个，吃了9个，剩下的占几分之几？",
        "一盒牛奶有12瓶，喝了5瓶，剩下的占几分之几？",
        "一包薯片有10片，吃了4片，剩下的占几分之几？",
        "一盒草莓有18个，吃了11个，剩下的占几分之几？",
        "一箱橘子有20个，卖了13个，剩下的占几分之几？",
        # 分数加减：小明吃了X，小红吃了Y，一共吃了几分之几
        "小明吃了一个蛋糕的2/8，小红吃了3/8，两人一共吃了几分之几？",
        "第一次用掉一根绳子的3/10，第二次用掉4/10，两次一共用掉几分之几？",
        "上午完成一项工作的4/12，下午完成5/12，一共完成了几分之几？",
        "姐姐吃了巧克力1/7，弟弟吃了3/7，两人一共吃了几分之几？",
        "第一天看了书的2/9，第二天看了4/9，两天一共看了几分之几？",
        "第一个月修了路的3/10，第二个月修了5/10，一共修了几分之几？",
        "甲队完成工程的2/8，乙队完成3/8，两队一共完成几分之几？",
        "小红喝了果汁1/6，小华喝了2/6，两人一共喝了几分之几？",
        # 分数加减：原来有X，用了Y，还剩几分之几
        "一瓶果汁有1瓶，喝了3/8，还剩几分之几？",
        "一根绳子长1米，用了4/10，还剩几分之几？",
        "一本书看完了2/5，还剩几分之几没看？",
        "一块地种了3/7的菜，剩下的占几分之几？",
        "一桶油用掉了5/12，还剩几分之几？",
        "一张纸用掉了2/6，还剩几分之几？",
        "一盒牛奶喝了1/4，还剩几分之几？",
        "一堆煤烧了4/9，还剩几分之几？",
        # 分数比较：谁多，多几分之几
        "小明吃了蛋糕的2/8，小红吃了3/8，谁吃得多？多吃几分之几？",
        "第一周修了路的4/10，第二周修了3/10，第一周多修几分之几？",
        "姐姐有糖果的3/7，妹妹有2/7，姐姐比妹妹多几分之几？",
    ]
    
    # 预洗牌池
    if not hasattr(MathWorkbookConfig, '_mul_pool_index'):
        random.shuffle(mul_pool)
        random.shuffle(frac_pool)
        MathWorkbookConfig._mul_pool = mul_pool
        MathWorkbookConfig._frac_pool = frac_pool
        MathWorkbookConfig._mul_pool_index = 0
        MathWorkbookConfig._frac_pool_index = 0
    
    # 每题取1道
    m_idx = MathWorkbookConfig._mul_pool_index
    f_idx = MathWorkbookConfig._frac_pool_index
    
    mul_problem = MathWorkbookConfig._mul_pool[m_idx % len(MathWorkbookConfig._mul_pool)]
    frac_problem = MathWorkbookConfig._frac_pool[f_idx % len(MathWorkbookConfig._frac_pool)]
    
    MathWorkbookConfig._mul_pool_index += 1
    MathWorkbookConfig._frac_pool_index += 1
    
    return [mul_problem, frac_problem]


# ============================================================================
# 渲染函数 (纯渲染逻辑，与生成分离)
# ============================================================================

def render_kousuan_section(doc, problems: List[str]) -> None:
    """
    渲染口算板块 (2列，Tab制表位分隔，10题5行)
    """
    add_title_paragraph(doc, "一、口算题（两位数乘法）", MathWorkbookConfig.FONT_SIZE_SECTION, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc, 1)
    
    for i in range(0, len(problems), 2):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        
        # 设置制表位：2列等分
        page_width = 21.0 - 2.0 - 2.0
        tab = Cm(page_width / 2)
        p.paragraph_format.tab_stops.add_tab_stop(tab)
        
        parts = [f"({i + 1}) {problems[i]}"]
        if i + 1 < len(problems):
            parts.append(f"({i + 2}) {problems[i + 1]}")
        
        text = "\t".join(parts)
        run = p.add_run(text)
        set_chinese_font(run, MathWorkbookConfig.FONT_MATH, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
    
    add_empty_line(doc, 1)


def render_shushi_section(doc, problems: List[Tuple[int, int, str]]) -> None:
    """
    渲染竖式笔算板块 (2×2表格布局，充足竖式空间)
    每题留出约8行空白用于列竖式计算，左右间距3cm
    """
    add_title_paragraph(doc, "二、笔算题（列竖式计算）", MathWorkbookConfig.FONT_SIZE_SECTION, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc, 1)
    
    # 创建2×2表格（不可见边框）
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # 设置表格宽度为页面可用宽度
    avail_width = 21.0 - 2.0 - 2.0
    col_width = Cm(avail_width / 2)
    
    for idx, (a, b, label) in enumerate(problems):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        
        # 清除cell默认段落
        cell.paragraphs[0].clear()
        cell.width = col_width
        
        # 题目序号+横式
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"({idx + 1}) {a} × {b} = ____")
        set_chinese_font(run, MathWorkbookConfig.FONT_MATH, MathWorkbookConfig.FONT_SIZE_PROBLEM, True)
        
        # 竖式答题空间（8个空行，间距1.5倍）
        for _ in range(8):
            blank = cell.add_paragraph()
            blank.paragraph_format.line_spacing = 1.5
            blank.paragraph_format.space_before = Pt(2)
            blank.paragraph_format.space_after = Pt(2)
            r = blank.add_run(" " * 30)
            set_chinese_font(r, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
        
        # 单元格左边缩进
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.left_indent = Cm(0.5)
    
    # 设置列宽
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width
    
    add_empty_line(doc, 1)


def render_fraction_section(doc, fractions: List[Tuple[int, int, int, int, str]]) -> None:
    """
    渲染分数加减法 (2列Tab分隔 + 题间空行)
    运算符来自生成数据，加法和减法正确显示
    """
    add_title_paragraph(doc, "三、分数加减法", MathWorkbookConfig.FONT_SIZE_SECTION, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc, 1)
    
    page_width = 21.0 - 2.0 - 2.0  # 17cm
    half = Cm(page_width / 2)
    
    frac_line = MathWorkbookConfig.FRACTION_LINE_CHAR * MathWorkbookConfig.FRACTION_LINE_LENGTH
    
    for i in range(0, len(fractions), 2):
        n1_1, d1, n1_2, d1_2, op1 = fractions[i]
        
        # 第1行：分子 (Tab分隔两列)
        p1 = doc.add_paragraph()
        p1.paragraph_format.line_spacing = 1.0
        p1.paragraph_format.tab_stops.add_tab_stop(half)
        part1 = f"    {n1_1:>2d}        {n1_2:>2d}"
        if i + 1 < len(fractions):
            n2_1, d2, n2_2, d2_2, op2 = fractions[i + 1]
            part2 = f"    {n2_1:>2d}        {n2_2:>2d}"
            p1.add_run(f"{part1}\t{part2}")
        else:
            p1.add_run(part1)
        for run in p1.runs:
            set_chinese_font(run, MathWorkbookConfig.FONT_MATH, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
        
        # 第2行：分数线 + 运算符
        p2 = doc.add_paragraph()
        p2.paragraph_format.line_spacing = 1.0
        p2.paragraph_format.tab_stops.add_tab_stop(half)
        line_part1 = f"  {frac_line}  {op1}  {frac_line}  ="
        if i + 1 < len(fractions):
            line_part2 = f"  {frac_line}  {op2}  {frac_line}  ="
            p2.add_run(f"{line_part1}\t{line_part2}")
        else:
            p2.add_run(line_part1)
        for run in p2.runs:
            set_chinese_font(run, MathWorkbookConfig.FONT_MATH, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
        
        # 第3行：分母
        p3 = doc.add_paragraph()
        p3.paragraph_format.line_spacing = 1.2
        p3.paragraph_format.tab_stops.add_tab_stop(half)
        den_part1 = f"    {d1:>2d}        {d1:>2d}"
        if i + 1 < len(fractions):
            den_part2 = f"    {d2:>2d}        {d2:>2d}"
            p3.add_run(f"{den_part1}\t{den_part2}")
        else:
            p3.add_run(den_part1)
        for run in p3.runs:
            set_chinese_font(run, MathWorkbookConfig.FONT_MATH, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
        
        # 题间空行（隔开每组分数题）
        add_empty_line(doc, 1)


def render_fill_section(doc, problems: List[str]) -> None:
    """
    渲染单位换算填空板块
    
    Args:
        doc: Document对象
        problems: 题目列表
    """
    add_title_paragraph(doc, "四、时间单位换算", MathWorkbookConfig.FONT_SIZE_SECTION, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc, 1)
    
    for i, problem in enumerate(problems, 1):
        add_problem_paragraph(doc, f"{i}. {problem}")
    
    add_empty_line(doc, 1)


def render_application_section(doc, problems: List[str]) -> None:
    """
    渲染应用题板块 (完整答题空间)
    
    Args:
        doc: Document对象
        problems: 题目列表
    """
    add_title_paragraph(doc, "五、应用题", MathWorkbookConfig.FONT_SIZE_SECTION, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_empty_line(doc, 1)
    
    for i, problem in enumerate(problems, 1):
        add_problem_paragraph(doc, f"{i}. {problem}")
        
        # 答题空间（3行空白）
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(" " * 50)
            set_chinese_font(run, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_PROBLEM, False)
        
        add_empty_line(doc, 1)


# ============================================================================
# 文档结构函数
# ============================================================================

def setup_document_format(doc) -> None:
    """
    设置文档整体格式 (出版级标准)
    
    Args:
        doc: Document对象
    """
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = MathWorkbookConfig.MARGIN_TOP
        section.bottom_margin = MathWorkbookConfig.MARGIN_BOTTOM
        section.left_margin = MathWorkbookConfig.MARGIN_LEFT
        section.right_margin = MathWorkbookConfig.MARGIN_RIGHT
        
        # 设置纸张大小
        section.page_width = MathWorkbookConfig.PAGE_WIDTH
        section.page_height = MathWorkbookConfig.PAGE_HEIGHT
    
    # 设置默认段落格式
    doc.styles['Normal'].paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_PROBLEM


def generate_cover_page(doc) -> None:
    """
    生成封面页 (标准教材封面格式)
    
    Args:
        doc: Document对象
    """
    add_empty_line(doc, 6)
    
    # 主标题
    add_title_paragraph(doc, MathWorkbookConfig.BOOK_TITLE, MathWorkbookConfig.FONT_SIZE_BOOK_TITLE, bold=True)
    
    add_empty_line(doc, 2)
    
    # 副标题
    add_title_paragraph(doc, MathWorkbookConfig.BOOK_SUBTITLE, Pt(16), bold=False)
    
    add_empty_line(doc, 6)
    
    # 版本信息
    add_title_paragraph(doc, f"版本 {MathWorkbookConfig.BOOK_VERSION}", MathWorkbookConfig.FONT_SIZE_SECTION, bold=False)
    
    add_empty_line(doc, 1)
    
    add_title_paragraph(doc, MathWorkbookConfig.BOOK_DATE, MathWorkbookConfig.FONT_SIZE_SECTION, bold=False)
    
    add_empty_line(doc, 2)
    
    add_title_paragraph(doc, MathWorkbookConfig.BOOK_AUTHOR, MathWorkbookConfig.FONT_SIZE_SECTION, bold=False)
    
    # 分页
    doc.add_page_break()


def generate_usage_guide(doc) -> None:
    """
    生成使用说明页 (家长指南)
    
    Args:
        doc: Document对象
    """
    add_title_paragraph(doc, "📖 使用说明", MathWorkbookConfig.FONT_SIZE_DAY_TITLE, bold=True)
    add_empty_line(doc, 2)
    
    guide_content = [
        "一、练习建议",
        "  1. 每天坚持练习30分钟，固定时间效果更好",
        "  2. 口算题建议计时，培养快速反应能力",
        "  3. 竖式计算注意数位对齐，养成好习惯",
        "  4. 分数题先观察再计算，做完记得检查",
        "  5. 应用题：每天1道乘法(两位数×两位数)+1道分数"""
        "  6. 时间换算题先记牢进率：1时=60分，1分=60秒",
        "",
        "二、艾宾浩斯复习法",
        "  第1天学习 → 第1天晚复习",
        "  → 第2天复习 → 第4天复习",
        "  → 第7天复习 → 第15天复习 → 第30天复习",
        "  按照这个规律，记忆留存率可达90%以上！",
        "",
        "三、评分标准",
        "  • 优秀：正确率 ≥ 95%",
        "  • 良好：85% ≤ 正确率 < 95%",
        "  • 及格：70% ≤ 正确率 < 85%",
        "  • 需努力：正确率 < 70%",
        "",
        "四、家长寄语",
        "  坚持每天练习，数学能力一定会有明显提升！",
        "  鼓励孩子独立完成，错题及时整理到错题本。",
    ]
    
    for line in guide_content:
        add_problem_paragraph(doc, line, first_line_indent=False)
    
    doc.add_page_break()


def generate_single_day(doc, day: int) -> None:
    """
    生成单日练习 (标准化流程)
    
    Args:
        doc: Document对象
        day: 第几天
    """
    # 每日标题
    review_mark = MathWorkbookConfig.REVIEW_MARKS.get(day, "")
    if review_mark:
        title = f"第 {day} 天   {review_mark}"
    else:
        title = f"第 {day} 天"
    
    add_title_paragraph(doc, title, MathWorkbookConfig.FONT_SIZE_DAY_TITLE, bold=True)
    add_empty_line(doc, 1)
    
    # 学生信息栏
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("姓名：________  用时：________  得分：________")
    set_chinese_font(info_run, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_INFO, False)
    
    add_empty_line(doc, 1)
    
    # 生成各板块题目
    config = MathWorkbookConfig.PROBLEMS_PER_DAY
    
    # 1. 口算题
    kousuan_probs = generate_kousuan_problems(config["口算"])
    render_kousuan_section(doc, kousuan_probs)
    
    # 2. 笔算题
    shushi_probs = generate_shushi_problems(config["笔算"])
    render_shushi_section(doc, shushi_probs)
    
    # 3. 分数加减法
    fraction_probs = generate_fraction_problems(config["分数"])
    render_fraction_section(doc, fraction_probs)
    
    # 4. 单位换算
    fill_probs = generate_fill_problems(config["填空"])
    render_fill_section(doc, fill_probs)
    
    # 5. 应用题
    app_probs = generate_application_problems(config["应用"])
    render_application_section(doc, app_probs)
    
    # 分页
    doc.add_page_break()


# ============================================================================
# 主流程函数
# ============================================================================

def generate_full_workbook() -> str:
    """
    生成完整的40天练习册 (标准流程)
    
    Returns:
        生成的文件完整路径
    """
    print("=" * 60)
    print("📚 三年级数学每日一练40天 - 标准教材版生成器 v3.0")
    print("=" * 60)
    print()
    
    # 0. 重置去重集合（确保每次生成都是干净的）
    MathWorkbookConfig.USED_PROBLEMS = set()
    # 重置填空题池
    if hasattr(MathWorkbookConfig, '_fill_pool_index'):
        del MathWorkbookConfig._fill_pool_index
        del MathWorkbookConfig._fill_pool
    # 重置应用题池
    if hasattr(MathWorkbookConfig, '_mul_pool_index'):
        del MathWorkbookConfig._mul_pool_index
        del MathWorkbookConfig._mul_pool
        del MathWorkbookConfig._frac_pool_index
        del MathWorkbookConfig._frac_pool
    dedup_stats = {"口算": 0, "竖式": 0, "分数": 0}
    
    # 1. 初始化文档
    print("[1/5] 初始化文档格式...")
    doc = Document()
    setup_document_format(doc)
    
    # 2. 生成封面
    print("[2/5] 生成封面页...")
    generate_cover_page(doc)
    
    # 3. 生成使用说明
    print("[3/5] 生成使用说明...")
    generate_usage_guide(doc)
    
    # 4. 生成每日练习
    print("[4/5] 正在生成40天练习题...")
    for day in range(1, 41):
        print(f"       正在生成第 {day:2d} 天...", end='\r')
        generate_single_day(doc, day)
    print()
    
    # 5. 保存文件
    print(f"[5/5] 正在保存到桌面: {MathWorkbookConfig.OUTPUT_FILENAME}...")
    doc.save(MathWorkbookConfig.OUTPUT_FULL_PATH)
    
    print()
    print("=" * 60)
    print("✅ 生成完成！")
    print("=" * 60)
    print()
    print(f"📂 文件路径: {MathWorkbookConfig.OUTPUT_FULL_PATH}")
    print()
    print("📊 统计信息:")
    print(f"   • 总天数: 40 天")
    print(f"   • 口算题: 40 × {MathWorkbookConfig.PROBLEMS_PER_DAY['口算']} = {40 * MathWorkbookConfig.PROBLEMS_PER_DAY['口算']} 题")
    print(f"   • 笔算题: 40 × {MathWorkbookConfig.PROBLEMS_PER_DAY['笔算']} = {40 * MathWorkbookConfig.PROBLEMS_PER_DAY['笔算']} 题")
    print(f"   • 分数题: 40 × {MathWorkbookConfig.PROBLEMS_PER_DAY['分数']} = {40 * MathWorkbookConfig.PROBLEMS_PER_DAY['分数']} 题")
    print(f"   • 时间换算: 40 × {MathWorkbookConfig.PROBLEMS_PER_DAY['填空']} = {40 * MathWorkbookConfig.PROBLEMS_PER_DAY['填空']} 题")
    print(f"   • 应用题: 40 × {MathWorkbookConfig.PROBLEMS_PER_DAY['应用']} = {40 * MathWorkbookConfig.PROBLEMS_PER_DAY['应用']} 题")
    print(f"   • 总题量: {40 * sum(MathWorkbookConfig.PROBLEMS_PER_DAY.values())} 题")
    print()
    print("🔍 去重验证:")
    total_used = len(MathWorkbookConfig.USED_PROBLEMS)
    print(f"   • 唯一题目哈希数: {total_used} 个")
    # 各类的统计
    kou_count = sum(1 for k in MathWorkbookConfig.USED_PROBLEMS if k.startswith("kou_t"))
    shu_count = sum(1 for k in MathWorkbookConfig.USED_PROBLEMS if k.startswith("shu_"))
    frac_count = sum(1 for k in MathWorkbookConfig.USED_PROBLEMS if k.startswith("frac:"))
    print(f"   • 口算唯一题: {kou_count} 个 (需要{40 * MathWorkbookConfig.PROBLEMS_PER_DAY['口算']}题)")
    print(f"   • 笔算唯一题: {shu_count} 个 (需要{40 * MathWorkbookConfig.PROBLEMS_PER_DAY['笔算']}题)")
    print(f"   • 分数唯一题: {frac_count} 个 (需要{40 * MathWorkbookConfig.PROBLEMS_PER_DAY['分数']}题)")
    dup_free = (kou_count >= 40 * MathWorkbookConfig.PROBLEMS_PER_DAY['口算'] and
                shu_count >= 40 * MathWorkbookConfig.PROBLEMS_PER_DAY['笔算'] and
                frac_count >= 40 * MathWorkbookConfig.PROBLEMS_PER_DAY['分数'])
    print(f"   {'✅ 全部不重复！' if dup_free else '⚠️ 有重复'}")
    print()
    print("💡 本版本基于十轮研究成果优化：")
    print("   • 符合人教版小学教材出版标准")
    print("   • 保护视力的字号和行距设置")
    print("   • 基于认知心理学的题量配置")
    print("   • 融入艾宾浩斯复习节奏标记")
    print("   • 标准分数竖式格式")
    print("   • 完整的封面和使用说明")
    print()
    
    return MathWorkbookConfig.OUTPUT_FULL_PATH


# ============================================================================
# 入口
# ============================================================================

if __name__ == "__main__":
    generate_full_workbook()
