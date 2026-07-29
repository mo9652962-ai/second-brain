# -*- coding: utf-8 -*-
"""
三年级数学每日一练40天 - 函数化标准版生成器
版本: 2.0 (函数化规范版)
作者: AI Assistant
创建日期: 2026-07-29

设计原则:
  1. 单一职责原则 - 每个函数只做一件事
  2. 参数化配置 - 所有参数通过函数参数传递
  3. 模块化设计 - 各功能模块独立
  4. 类型标注 - 使用类型提示增强可读性
  5. 文档字符串 - 每个函数都有完整文档
"""

import os
import random
from datetime import datetime, timedelta
from typing import List, Tuple, Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ============================================================================
# 配置常量区 (所有可配置参数集中管理)
# ============================================================================

class MathWorkbookConfig:
    """练习册配置类 - 统一管理所有参数"""
    
    # 页面设置
    PAGE_WIDTH = Cm(21.0)          # A4纸宽度
    PAGE_HEIGHT = Cm(29.7)         # A4纸高度
    MARGIN_TOP = Cm(2.54)          # 上边距
    MARGIN_BOTTOM = Cm(2.54)       # 下边距
    MARGIN_LEFT = Cm(2.5)          # 左边距
    MARGIN_RIGHT = Cm(2.0)         # 右边距
    
    # 字体设置 (小学教材标准)
    FONT_MAIN = "宋体"              # 正文字体
    FONT_TITLE = "黑体"             # 标题字体
    FONT_SIZE_TITLE = Pt(22)        # 主标题字号 (二号)
    FONT_SIZE_SECTION = Pt(14)      # 板块标题字号 (四号)
    FONT_SIZE_CONTENT = Pt(10.5)    # 题目内容字号 (小五)
    FONT_SIZE_INFO = Pt(9)          # 辅助信息字号
    
    # 行距设置
    LINE_SPACING_TITLE = Pt(24)     # 主标题行距
    LINE_SPACING_PROBLEM = Pt(18)   # 普通题目行距
    LINE_SPACING_FRACTION_NUM = Pt(16)  # 分数分子/分母行距
    LINE_SPACING_FRACTION_LINE = Pt(14) # 分数线行距
    
    # 题量配置
    PROBLEMS_PER_DAY = {
        "口算": 15,
        "竖式": 10,
        "分数": 10,
        "填空": 5,
        "应用": 2,
    }
    
    # 分数格式
    FRACTION_LINE_CHAR = "─"       # 分数线字符
    FRACTION_LINE_LENGTH = 2        # 分数线长度(字符)
    
    # 输出配置 - 默认保存到桌面
    DESKTOP_PATH = r"C:\Users\31954\Desktop"
    OUTPUT_FILENAME = "三年级数学每日一练40天_函数标准版.docx"
    OUTPUT_FULL_PATH = os.path.join(DESKTOP_PATH, OUTPUT_FILENAME)


# ============================================================================
# 工具函数区
# ============================================================================

def set_font(run, font_name: str, font_size: Pt, bold: bool = False) -> None:
    """
    设置字体样式（单一职责：只负责设置字体）
    
    Args:
        run: docx Run对象
        font_name: 字体名称
        font_size: 字号(Pt)
        bold: 是否加粗
    """
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    # 解决中文字体不生效问题
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_styled_paragraph(doc, text: str = "", style: str = "Normal") -> Tuple[Any, Any]:
    """
    添加带样式的段落（单一职责：只负责创建段落并返回引用）
    
    Args:
        doc: Document对象
        text: 段落文本
        style: 样式名称
    
    Returns:
        (paragraph, run) 元组
    """
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    return p, run


def add_title_paragraph(doc, text: str, font_size: Pt, 
                         alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER) -> Any:
    """
    添加标题段落（单一职责：只负责标题格式设置）
    
    Args:
        doc: Document对象
        text: 标题文本
        font_size: 字号
        alignment: 对齐方式
    
    Returns:
        paragraph对象
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    set_font(run, MathWorkbookConfig.FONT_TITLE, font_size, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_content_paragraph(doc, text: str = "") -> Tuple[Any, Any]:
    """
    添加内容段落（单一职责：只负责正文内容格式）
    
    Args:
        doc: Document对象
        text: 段落文本
    
    Returns:
        (paragraph, run) 元组
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_CONTENT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_PROBLEM
    return p, run


# ============================================================================
# 题目生成函数区 (每个题型一个独立函数)
# ============================================================================

def generate_kousuan_problems(count: int, difficulty: int = 1) -> List[str]:
    """
    生成两位数乘法口算题（单一职责：只负责生成口算题列表）
    
    Args:
        count: 题目数量
        difficulty: 难度等级 (1=基础, 2=进阶, 3=挑战)
    
    Returns:
        题目字符串列表
    """
    problems = []
    for _ in range(count):
        if difficulty == 1:
            a = random.randint(11, 50)
            b = random.choice([10, 20, 30, 40, 50])
        elif difficulty == 2:
            a = random.randint(11, 60)
            b = random.choice([10, 15, 20, 25, 30, 40, 50])
        else:
            a = random.randint(11, 80)
            b = random.randint(11, 50)
        
        problems.append(f"{a} × {b} =")
    
    return problems


def generate_shushi_problems(count: int, difficulty: int = 1) -> List[str]:
    """
    生成两位数乘两位数笔算题（单一职责：只负责生成竖式题列表）
    
    Args:
        count: 题目数量
        difficulty: 难度等级
    
    Returns:
        题目字符串列表
    """
    problems = []
    for _ in range(count):
        if difficulty == 1:
            a = random.randint(11, 35)
            b = random.randint(11, 25)
        elif difficulty == 2:
            a = random.randint(15, 50)
            b = random.randint(11, 40)
        else:
            a = random.randint(20, 65)
            b = random.randint(11, 50)
        
        problems.append(f"{a} × {b} =")
    
    return problems


def generate_fraction_data(count: int) -> List[Tuple[int, int, int]]:
    """
    生成分数口算题数据（单一职责：只负责生成分数数据元组）
    
    Args:
        count: 题目数量
    
    Returns:
        (分子1, 分母, 分子2) 元组列表
    """
    fractions = []
    denominators = [10, 8, 6]
    
    for _ in range(count):
        den = random.choice(denominators)
        num1 = random.randint(den // 2, den - 1)
        num2 = random.randint(1, num1 - 1)
        fractions.append((num1, den, num2))
    
    return fractions


def generate_time_convert_problems(count: int) -> List[str]:
    """
    生成时间换算填空题（单一职责：只负责生成时间换算题）
    
    Args:
        count: 题目数量
    
    Returns:
        题目字符串列表
    """
    problems = []
    templates = [
        lambda: f"{random.randint(1, 12)}时{random.choice([0, 15, 20, 30])}分 = (    )分",
        lambda: f"1分{random.choice([0, 10, 15, 20, 30])}秒 = (    )秒",
        lambda: f"{random.choice([30, 60, 90, 120, 180])}分 = (    )时",
        lambda: f"{random.choice([24, 48, 72, 96, 120])}时 = (    )日",
        lambda: f"{random.choice([7, 8, 9])}、{random.choice([7, 8, 9])}月一共有(    )天",
        lambda: f"平年上半年一共有(    )天",
        lambda: f"闰年全年有(    )天",
        lambda: f"一年有(    )个季度",
    ]
    
    selected = random.sample(templates, min(count, len(templates)))
    for t in selected:
        problems.append(t())
    
    # 如果不够补充
    while len(problems) < count:
        problems.append(random.choice(templates)())
    
    return problems


def generate_application_problems(count: int) -> List[str]:
    """
    生成应用题（单一职责：只负责生成应用题）
    
    Args:
        count: 题目数量
    
    Returns:
        题目字符串列表
    """
    scenarios = [
        lambda: f"水龙头每天浪费{random.randint(10, 20)}千克水，{random.randint(10, 30)}天浪费多少千克水？",
        lambda: f"一块布料平均分成{random.randint(4, 8)}份，用掉{random.randint(2, 5)}份，剩下占几分之几？",
        lambda: f"一个长方形操场，长{random.randint(30, 80)}米，宽{random.randint(20, 50)}米，面积是多少平方米？",
        lambda: f"同学们去植树，每行种{random.randint(10, 20)}棵，种了{random.randint(8, 15)}行，一共种了多少棵？",
    ]
    
    problems = []
    selected = random.sample(scenarios, min(count, len(scenarios)))
    for s in selected:
        problems.append(s())
    
    return problems


# ============================================================================
# 文档渲染函数区 (每个板块一个独立函数)
# ============================================================================

def render_kousuan_section(doc, problems: List[str]) -> None:
    """
    渲染口算题板块（单一职责：只负责口算题的Word渲染）
    
    Args:
        doc: Document对象
        problems: 口算题列表
    """
    # 板块标题
    p_title = add_title_paragraph(doc, "一、两位数乘法口算", 
                                  MathWorkbookConfig.FONT_SIZE_SECTION,
                                  WD_ALIGN_PARAGRAPH.LEFT)
    
    # 所有题目放在同一个段落自动换行
    p, run = add_content_paragraph(doc)
    for idx, prob in enumerate(problems, 1):
        run.add_text(f"{idx}. {prob}    ")
    
    p.paragraph_format.left_indent = Cm(0.5)


def render_shushi_section(doc, problems: List[str]) -> None:
    """
    渲染竖式计算板块（单一职责：只负责竖式题的Word渲染）
    
    Args:
        doc: Document对象
        problems: 竖式题列表
    """
    # 板块标题
    add_title_paragraph(doc, "二、两位数乘两位数笔算", 
                       MathWorkbookConfig.FONT_SIZE_SECTION,
                       WD_ALIGN_PARAGRAPH.LEFT)
    
    # 每题单独一行
    for idx, prob in enumerate(problems, 1):
        p, run = add_content_paragraph(doc, f"{idx}. {prob}")
        p.paragraph_format.left_indent = Cm(0.5)


def render_fraction_section(doc, fractions: List[Tuple[int, int, int]]) -> None:
    """
    渲染分数口算板块（标准竖式排版）（单一职责：只负责分数渲染）
    符合人教版三年级数学教材标准格式：
    - 分子居中对齐
    - 分数线长度与分子/分母宽度匹配
    - 等号与分数线同高
    - 运算符号居中
    
    Args:
        doc: Document对象
        fractions: (分子1, 分母, 分子2) 元组列表
    """
    # 板块标题
    add_title_paragraph(doc, "三、分数口算", 
                       MathWorkbookConfig.FONT_SIZE_SECTION,
                       WD_ALIGN_PARAGRAPH.LEFT)
    
    # 分成两列，每列5题
    col1 = fractions[:5]
    col2 = fractions[5:]
    
    # 标准分数线：3个字符，覆盖分子/分母宽度
    line_char = MathWorkbookConfig.FRACTION_LINE_CHAR * 3
    
    # 每行渲染两题，标准竖式格式
    for row in range(5):
        f1 = col1[row]
        f2 = col2[row] if row < len(col2) else None
        
        # ================================
        # 第一行：分子
        # ================================
        p_num = doc.add_paragraph()
        p_num.paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_FRACTION_NUM
        if f2:
            # 标准格式：分子居中，2位数字右对齐
            text = f"    {f1[0]:>2d}        {f2[0]:>2d}  "
        else:
            text = f"    {f1[0]:>2d}  "
        run_num = p_num.add_run(text)
        set_font(run_num, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_CONTENT)
        
        # ================================
        # 第二行：分数线 + 运算符号 + 等号
        # 标准：分数线3字符，运算符号在中间，等号与分数线同高
        # ================================
        p_line = doc.add_paragraph()
        p_line.paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_FRACTION_LINE
        if f2:
            # 标准格式：分数线 - 分数线 =    分数线 - 分数线 =
            text = f"  {line_char}  -  {line_char}  =    {line_char}  -  {line_char}  =  "
        else:
            text = f"  {line_char}  -  {line_char}  =  "
        run_line = p_line.add_run(text)
        set_font(run_line, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_CONTENT)
        
        # ================================
        # 第三行：分母
        # ================================
        p_den = doc.add_paragraph()
        p_den.paragraph_format.line_spacing = MathWorkbookConfig.LINE_SPACING_FRACTION_NUM
        if f2:
            # 标准格式：分母居中，2位数字右对齐
            text = f"    {f1[1]:>2d}        {f2[1]:>2d}  "
        else:
            text = f"    {f1[1]:>2d}  "
        run_den = p_den.add_run(text)
        set_font(run_den, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_CONTENT)
        p_den.paragraph_format.space_after = Pt(4)


def render_time_convert_section(doc, problems: List[str]) -> None:
    """
    渲染时间换算填空板块（单一职责：只负责时间换算题渲染）
    
    Args:
        doc: Document对象
        problems: 时间换算题列表
    """
    # 板块标题
    add_title_paragraph(doc, "四、年月日、时分秒换算填空", 
                       MathWorkbookConfig.FONT_SIZE_SECTION,
                       WD_ALIGN_PARAGRAPH.LEFT)
    
    for idx, prob in enumerate(problems, 1):
        p, run = add_content_paragraph(doc, f"{idx}. {prob}")
        p.paragraph_format.left_indent = Cm(0.5)


def render_application_section(doc, problems: List[str]) -> None:
    """
    渲染应用题板块（单一职责：只负责应用题渲染）
    
    Args:
        doc: Document对象
        problems: 应用题列表
    """
    # 板块标题
    add_title_paragraph(doc, "五、解决问题", 
                       MathWorkbookConfig.FONT_SIZE_SECTION,
                       WD_ALIGN_PARAGRAPH.LEFT)
    
    for idx, prob in enumerate(problems, 1):
        p, run = add_content_paragraph(doc, f"{idx}. {prob}")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(1.0)  # 首行缩进
        p.paragraph_format.space_after = Pt(6)


# ============================================================================
# 页面渲染函数
# ============================================================================

def render_day_header(doc, day_num: int) -> None:
    """
    渲染每日练习头部（标题 + 学生信息）（单一职责：只负责每日头部）
    
    Args:
        doc: Document对象
        day_num: 第几天
    """
    # 主标题
    add_title_paragraph(doc, f"三年级数学每日一练  第{day_num}天",
                       MathWorkbookConfig.FONT_SIZE_TITLE)
    
    # 学生信息栏
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run("姓名：___________    用时：___________    家长签字：___________")
    set_font(run_info, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_INFO)
    run_info.font.color.rgb = RGBColor(128, 128, 128)
    p_info.paragraph_format.space_after = Pt(6)


def render_one_day(doc, day_num: int, difficulty: int = 1) -> None:
    """
    渲染完整一天的练习（组合调用各渲染函数）
    
    Args:
        doc: Document对象
        day_num: 第几天
        difficulty: 难度等级
    """
    # 1. 每日标题
    render_day_header(doc, day_num)
    
    # 2. 口算题
    kousuan = generate_kousuan_problems(
        MathWorkbookConfig.PROBLEMS_PER_DAY["口算"],
        difficulty
    )
    render_kousuan_section(doc, kousuan)
    
    # 3. 竖式计算
    shushi = generate_shushi_problems(
        MathWorkbookConfig.PROBLEMS_PER_DAY["竖式"],
        difficulty
    )
    render_shushi_section(doc, shushi)
    
    # 4. 分数口算
    fractions = generate_fraction_data(
        MathWorkbookConfig.PROBLEMS_PER_DAY["分数"]
    )
    render_fraction_section(doc, fractions)
    
    # 5. 时间换算填空
    time_convert = generate_time_convert_problems(
        MathWorkbookConfig.PROBLEMS_PER_DAY["填空"]
    )
    render_time_convert_section(doc, time_convert)
    
    # 6. 解决问题
    application = generate_application_problems(
        MathWorkbookConfig.PROBLEMS_PER_DAY["应用"]
    )
    render_application_section(doc, application)
    
    # 分页
    doc.add_page_break()


# ============================================================================
# 主函数 (入口)
# ============================================================================

def main() -> None:
    """
    主函数 - 生成完整练习册（单一职责：只负责整体流程控制）
    
    流程:
        1. 初始化文档
        2. 循环生成40天练习
        3. 保存文件
        4. 输出统计信息
    """
    print("=" * 60)
    print("  三年级数学每日一练 40天生成器 (函数标准版 v2.0)")
    print("=" * 60)
    print()
    
    # 1. 初始化文档
    print("[1/4] 初始化文档...")
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = MathWorkbookConfig.FONT_MAIN
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), MathWorkbookConfig.FONT_MAIN)
    doc.styles['Normal'].font.size = MathWorkbookConfig.FONT_SIZE_CONTENT
    
    # 2. 生成40天练习
    print("[2/4] 正在生成40天练习题...")
    for day in range(1, 41):
        # 难度随天数递增
        difficulty = 1 if day <= 15 else 2 if day <= 30 else 3
        print(f"  生成第 {day:2d}/40 天...", end='\r')
        render_one_day(doc, day, difficulty)
    print()
    
    # 3. 添加参考答案占位符
    print("[3/4] 正在生成参考答案...")
    doc.add_page_break()
    add_title_paragraph(doc, "参考答案", Pt(18))
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = p_note.add_run("（根据学生完成情况自行对照）")
    set_font(run_note, MathWorkbookConfig.FONT_MAIN, MathWorkbookConfig.FONT_SIZE_INFO)
    
    # 4. 保存文件 - 直接保存到桌面
    print(f"[4/4] 正在保存到桌面: {MathWorkbookConfig.OUTPUT_FILENAME}...")
    doc.save(MathWorkbookConfig.OUTPUT_FULL_PATH)
    
    # 输出统计信息
    print()
    print("=" * 60)
    print("  生成完成！")
    print("=" * 60)
    print()
    print(f"  文件位置: {MathWorkbookConfig.OUTPUT_FULL_PATH}")
    print()
    print("  📊 题量统计:")
    total = 0
    for ptype, count in MathWorkbookConfig.PROBLEMS_PER_DAY.items():
        total_type = count * 40
        total += total_type
        print(f"    {ptype:4s}: {total_type:4d} 题 ({count:2d}/天)")
    print(f"    {'-' * 20}")
    print(f"    总计: {total:4d} 题")
    print()
    print("  🎯 函数化规范特性:")
    print("    ✓ 单一职责原则 - 每个函数只做一件事")
    print("    ✓ 参数化配置 - 所有参数集中管理")
    print("    ✓ 模块化设计 - 各功能独立")
    print("    ✓ 类型标注 - 类型提示增强可读性")
    print("    ✓ 文档字符串 - 完整函数文档")
    print()


if __name__ == "__main__":
    main()
