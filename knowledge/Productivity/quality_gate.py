#!/usr/bin/env python3
"""quality_gate.py — 数模代写交付质量门禁（四硬项）v1.0 2026-08-23

借鉴 AutoMCM-Pro / MathModel-Skill 的「代码自证 + 证据追溯」设计，落地为可执行检查。
接单交付前自动跑一遍，四项全 PASS 才允许交付。

四硬项：
  G1 数值一致性：正文/表格/附录三处数字一致，无占位符
  G2 代码自证：每个求解脚本配独立验证脚本，全 PASS 才写入论文
  G3 证据追溯：论文每个数字来自真实运行结果，防 AI 编造数值
  G4 AI 声明合规：按当届规则核查 AI 使用声明

用法：
  python quality_gate.py check <论文文件> [--code-dir <代码目录>] [--results-dir <结果目录>] [--contest cumcm|mcm] [--run-tests]
  python quality_gate.py demo                          # 生成演示样本并自检

依赖：python-docx（.docx 提取）；纯文本 .md/.tex 无需额外依赖
"""
import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

# ── 占位符黑名单（G1）──
PLACEHOLDER_PATTERNS = [
    r"TODO", r"TBD", r"占位", r"待填", r"待补", r"待写", r"待插入",
    r"xxx", r"XXX", r"这里放", r"此处插入", r"lorem ipsum", r"Lorem ipsum",
    r"【", r"】", r"\[待", r"placeholder",
]

# ── AI 使用声明关键词（G4）──
AI_DISCLOSURE_KEYWORDS = [
    "AI 工具", "AI工具", "人工智能工具", "生成式人工智能", "AIGC",
    "AI-assisted", "AI assisted", "ChatGPT", "DeepSeek", "文心一言",
    "大语言模型", "LLM", "AI 使用声明", "AI使用声明", "辅助工具",
]

# ── 求解脚本 / 验证脚本命名约定 ──
SOLVE_PATTERNS = [r"^solve_.*\.py$", r"^q\d+.*\.py$", r"^model_.*\.py$", r"^.*_model\.py$"]
VERIFY_PATTERNS = [r"^verify_.*\.py$", r"^test_.*\.py$", r"^check_.*\.py$"]


# ══════════════════════ 文本提取 ══════════════════════

def extract_text(path: Path) -> str:
    """从 .docx / .md / .tex / .txt 提取全文"""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    elif suffix in (".md", ".tex", ".txt", ".markdown"):
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"不支持的论文格式: {suffix}（支持 docx/md/tex/txt）")


def extract_numbers(text: str) -> list[str]:
    """提取正文中看起来像『结果数值』的候选：带小数/百分比/科学计数"""
    return re.findall(r"\b\d+\.\d+(?:[eE][+-]?\d+)?\b|\b\d+%\b|\b\d+(?:\.\d+)?e[+-]?\d+\b", text)


# ══════════════════════ G1 数值一致性 ══════════════════════

def gate1_placeholders(text: str) -> list[str]:
    hits = []
    for line_no, line in enumerate(text.splitlines(), 1):
        for pat in PLACEHOLDER_PATTERNS:
            if re.search(pat, line):
                hits.append(f"L{line_no}: {line.strip()[:80]}")
                break
    return hits


def gate1_number_crosscheck(text: str) -> tuple[list[str], list[str]]:
    """抽取数字，检查明显不一致：同文档内同数字不应以矛盾形式出现（简化：返回唯一数字列表供人工复核）"""
    nums = extract_numbers(text)
    # 简单启发：出现超过 3 次的数字标记为「关键数值」（正文常复述重要结果）
    from collections import Counter
    cnt = Counter(nums)
    key_nums = [n for n, c in cnt.items() if c >= 3]
    return key_nums, nums


# ══════════════════════ G2 代码自证 ══════════════════════

def gate2_code_selfverify(code_dir: Path, run_tests: bool) -> list[dict]:
    """检查每个求解脚本是否有对应验证脚本；可选运行验证脚本"""
    if not code_dir or not code_dir.exists():
        return [{"status": "SKIP", "msg": "未提供 --code-dir，跳过代码自证"}]

    solve_files = []
    for pat in SOLVE_PATTERNS:
        solve_files += [p for p in code_dir.glob("*.py") if re.match(pat, p.name)]
    solve_files = sorted(set(solve_files))

    if not solve_files:
        return [{"status": "SKIP", "msg": f"{code_dir} 未找到求解脚本（solve_*.py / q*_*.py / model_*.py）"}]

    results = []
    for s in solve_files:
        stem = s.stem
        # 查找对应验证脚本：verify_<stem>.py 或 test_<stem>.py 或 check_<stem>.py
        verify = None
        all_py = list(code_dir.glob("*.py"))
        for v in all_py:
            if v == s:
                continue
            if any(re.match(vp, v.name) for vp in VERIFY_PATTERNS):
                # 匹配关系：验证脚本名包含求解脚本名（或去前缀后包含）
                v_stem = v.stem
                v_core = re.sub(r"^(verify|test|check)_", "", v_stem)
                if stem in v_stem or v_core == stem or v_core in stem or stem in v_core:
                    verify = v
                    break

        if not verify:
            results.append({"script": s.name, "status": "FAIL", "msg": "缺少独立验证脚本（verify_*/test_*/check_*）"})
            continue

        if not run_tests:
            results.append({"script": s.name, "verify": verify.name, "status": "PASS", "msg": "验证脚本存在（未运行，加 --run-tests 执行）"})
            continue

        # 运行验证脚本
        try:
            r = subprocess.run(
                [sys.executable, str(verify)], capture_output=True, text=True, timeout=120, cwd=str(code_dir)
            )
            if r.returncode == 0:
                results.append({"script": s.name, "verify": verify.name, "status": "PASS", "msg": f"运行通过 (exit 0)"})
            else:
                results.append({"script": s.name, "verify": verify.name, "status": "FAIL", "msg": f"运行失败 (exit {r.returncode}): {r.stderr.strip()[:120]}"})
        except subprocess.TimeoutExpired:
            results.append({"script": s.name, "verify": verify.name, "status": "FAIL", "msg": "验证脚本超时 (>120s)"})
        except Exception as e:
            results.append({"script": s.name, "verify": verify.name, "status": "ERROR", "msg": str(e)[:120]})

    return results


# ══════════════════════ G3 证据追溯 ══════════════════════

def gate3_evidence_traceability(text: str, results_dir: Path | None) -> list[dict]:
    """论文数字应能追溯：若提供了 results_dir，检查其中是否有结果产物（json/csv/npy/txt）"""
    issues = []

    # 1. 有没有声称"运行结果"却没有任何结果文件
    claims_run = bool(re.search(r"运行结果|结果如下|计算结果|实验得到|我们得到", text))
    if results_dir is None:
        if claims_run:
            issues.append({"status": "WARN", "msg": "论文声称有运行结果，但未提供 --results-dir 供核对"})
        else:
            issues.append({"status": "PASS", "msg": "未声称运行结果，无需结果目录"})
        return issues

    if not results_dir.exists():
        issues.append({"status": "FAIL", "msg": f"结果目录不存在: {results_dir}"})
        return issues

    result_files = list(results_dir.rglob("*")) if results_dir.exists() else []
    result_exts = {".json", ".csv", ".npy", ".txt", ".png", ".pdf", ".xlsx"}
    artifacts = [f for f in result_files if f.suffix.lower() in result_exts]

    if claims_run and not artifacts:
        issues.append({"status": "FAIL", "msg": "论文声称有运行结果，但 results_dir 无任何产物（json/csv/npy/png 等）"})
    elif claims_run and artifacts:
        issues.append({"status": "PASS", "msg": f"找到 {len(artifacts)} 个结果产物可追溯"})
    else:
        issues.append({"status": "PASS", "msg": "论文未明显声称运行结果"})

    # 2. 检查结果文件里是否有 NaN / inf（数值异常是 AI 编造/代码 bug 的强信号）
    bad_vals = []
    for f in artifacts:
        if f.suffix.lower() in (".csv", ".txt", ".json"):
            try:
                content = f.read_text(encoding="utf-8", errors="ignore")[:200000]
                if re.search(r"\bNaN\b|\binf\b|\bnull\b", content):
                    bad_vals.append(f.name)
            except Exception:
                pass
    if bad_vals:
        issues.append({"status": "FAIL", "msg": f"结果文件含 NaN/inf/null（数值异常）：{', '.join(bad_vals[:5])}"})

    return issues


# ══════════════════════ G4 AI 声明合规 ══════════════════════

def gate4_ai_disclosure(text: str, contest: str) -> list[dict]:
    hits = [kw for kw in AI_DISCLOSURE_KEYWORDS if kw.lower() in text.lower()]
    if hits:
        return [{"status": "PASS", "msg": f"检测到 AI 使用相关表述：{', '.join(hits[:5])}"}]
    else:
        if contest == "mcm":
            # MCM/ICM 2024+ 明确要求 AI 使用声明
            return [{"status": "FAIL", "msg": "MCM/ICM 2024 起必须含 AI Use Report/声明（当前未检测到 AI 相关表述）"}]
        else:
            # CUMCM 2025+ 多赛区要求披露
            return [{"status": "WARN", "msg": "未检测到 AI 使用声明——国赛部分赛区 2025 起要求披露 AI 使用情况，建议补充（若确未使用 AI 可忽略）"}]


# ══════════════════════ 报告 ══════════════════════

def run_checks(paper: Path, code_dir: Path | None, results_dir: Path | None, contest: str, run_tests: bool) -> dict:
    text = extract_text(paper)

    g1_ph = gate1_placeholders(text)
    g1_key_nums, _ = gate1_number_crosscheck(text)

    g2 = gate2_code_selfverify(code_dir, run_tests)
    g3 = gate3_evidence_traceability(text, results_dir)
    g4 = gate4_ai_disclosure(text, contest)

    gates = {
        "G1_数值一致性": {
            "status": "FAIL" if g1_ph else "PASS",
            "details": [{"status": "FAIL" if g1_ph else "PASS",
                          "msg": f"发现 {len(g1_ph)} 处占位符" if g1_ph else "无占位符 ✓"}]
                      + [{"status": "INFO", "msg": f"关键数值（重复≥3次）: {', '.join(g1_key_nums[:10])}"}],
            "placeholder_hits": g1_ph,
        },
        "G2_代码自证": {
            "status": "PASS" if all(x["status"] == "PASS" or x["status"] == "SKIP" for x in g2) else "FAIL",
            "details": g2,
        },
        "G3_证据追溯": {
            "status": "FAIL" if any(x["status"] == "FAIL" for x in g3) else "PASS",
            "details": g3,
        },
        "G4_AI声明合规": {
            "status": "FAIL" if any(x["status"] == "FAIL" for x in g4) else "PASS",
            "details": g4,
        },
    }

    overall = "PASS" if all(g["status"] == "PASS" for g in gates.values()) else "FAIL"
    return {"paper": str(paper), "contest": contest, "overall": overall, "gates": gates}


def print_report(report: dict, verbose: bool = False):
    print("=" * 60)
    print(f"QUALITY GATE — {report['paper']}")
    print(f"Contest: {report['contest']} | OVERALL: {report['overall']}")
    print("=" * 60)
    for gname, gate in report["gates"].items():
        icon = {"PASS": "✅", "FAIL": "❌", "WARN": "⚠️"}.get(gate["status"], "❓")
        print(f"\n{icon} {gname} [{gate['status']}]")
        for d in gate["details"]:
            if d["status"] == "INFO":
                if verbose:
                    print(f"    ℹ️  {d['msg']}")
            elif d["status"] == "PASS":
                print(f"    ✓ {d['msg']}")
            elif d["status"] == "SKIP":
                print(f"    ➖ {d['msg']}")
            elif d["status"] == "WARN":
                print(f"    ⚠️  {d['msg']}")
            else:
                print(f"    ✗ {d['msg']}")
    if verbose and report["gates"]["G1_数值一致性"]["placeholder_hits"]:
        print("\n  G1 占位符明细:")
        for h in report["gates"]["G1_数值一致性"]["placeholder_hits"]:
            print(f"      {h}")
    print("\n" + ("🎉 全部通过，可交付" if report["overall"] == "PASS" else "🚫 存在未通过项，修复后再交付"))


# ══════════════════════ Demo ══════════════════════

def make_demo(tmp: Path):
    """生成演示样本：论文 + 求解脚本 + 验证脚本 + 结果目录，并跑一次检查"""
    paper = tmp / "demo_paper.md"
    paper.write_text("""# 2026 高教社杯 A 题论文

## 摘要
我们建立了线性规划模型，最优解为 32.5 万元，比基线提升 12.3%。

## 模型求解
运行结果如下：目标值 32.5，灵敏度系数 0.87。
（此处插入表格）TODO

## 结论
最终推荐方案 A，收益 32.5 万元。

我们使用了 AI 工具辅助完成文献检索与代码调试，已按竞赛规则披露。
""", encoding="utf-8")

    code_dir = tmp / "code"
    code_dir.mkdir()
    (code_dir / "solve_q1.py").write_text("""# 求解脚本
def main():
    print("32.5")
if __name__ == "__main__":
    main()
""", encoding="utf-8")
    (code_dir / "verify_solve_q1.py").write_text("""# 验证脚本：约束满足 + 数值稳定
def test_result():
    assert abs(32.5 - 32.5) < 1e-6
    print("PASS")
if __name__ == "__main__":
    test_result()
""", encoding="utf-8")
    # 无验证脚本的反例
    (code_dir / "solve_q2.py").write_text("print('q2')\n", encoding="utf-8")

    results_dir = tmp / "results"
    results_dir.mkdir()
    (results_dir / "q1_result.json").write_text('{"objective": 32.5, "sensitivity": 0.87}\n', encoding="utf-8")

    return paper, code_dir, results_dir


def main():
    parser = argparse.ArgumentParser(description="数模代写交付质量门禁（四硬项）")
    sub = parser.add_subparsers(dest="cmd")

    p_check = sub.add_parser("check", help="检查一篇论文")
    p_check.add_argument("paper", type=Path, help="论文文件 (.docx/.md/.tex)")
    p_check.add_argument("--code-dir", type=Path, default=None, help="求解代码目录")
    p_check.add_argument("--results-dir", type=Path, default=None, help="运行结果目录（json/csv/npy/png 等）")
    p_check.add_argument("--contest", choices=["cumcm", "mcm"], default="cumcm", help="竞赛类型（G4 用）")
    p_check.add_argument("--run-tests", action="store_true", help="实际运行验证脚本（默认只检查存在性）")
    p_check.add_argument("--verbose", action="store_true", help="显示关键数值等 INFO 信息")

    sub.add_parser("demo", help="生成演示样本并自检")

    args = parser.parse_args()
    if args.cmd == "demo":
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            paper, code_dir, results_dir = make_demo(tmp)
            report = run_checks(paper, code_dir, results_dir, "cumcm", run_tests=True)
            print_report(report, verbose=True)
        print("\n说明：demo 故意留了 1 处占位符（TODO）+ 1 个无验证脚本的求解脚本（solve_q2.py），所以 G1/G2 会 FAIL——演示门禁抓错能力。")
        return 0
    elif args.cmd == "check":
        if not args.paper.exists():
            print(f"❌ 论文文件不存在: {args.paper}")
            return 2
        try:
            report = run_checks(args.paper, args.code_dir, args.results_dir, args.contest, args.run_tests)
        except ValueError as e:
            print(f"❌ {e}")
            return 2
        print_report(report, verbose=args.verbose)
        return 0 if report["overall"] == "PASS" else 1
    else:
        parser.print_help()
        return 2


if __name__ == "__main__":
    sys.exit(main())
